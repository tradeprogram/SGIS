"""
공모전 제출물 생성 — 공식 양식 준수.

양식 근거 : 제6회 SGIS 활용 우수사례 공모전 개최 안내문(게시용) 붙임1 작성 양식.
  □ 추진배경 및 필요성 / □ 추진과정 / □ 추진내용 / □ 추진성과 4절 구성
  분량 5페이지 이내 서술형
  용지 A4, 여백 위·아래 10mm, 머리·꼬리말 10mm, 좌·우 20mm
  권장 서식 제목 HY헤드라인M 16pt, 본문 휴먼명조 15pt, 줄간격 160 — "서식변경 가능"
  추진 결과 근거자료는 별도 첨부(본문 5쪽에 포함되지 않는다)
  "제출물 양식을 준수하지 않으면 심사대상에서 제외될 수 있음"

  ※ 권장 15pt 를 그대로 쓰면 5쪽에 핵심이 안 들어간다. 양식이 서식변경을
    허용하므로 11pt·줄간격 1.3 으로 낮추되, 읽기 어려울 만큼 줄이지는 않는다.

심사 배점(전문가 60%) : 충실성 30 · 효과성 30 · 확산가능성 30 · 창의성 10
  → 추진내용을 통째로 "SGIS 활용"에 할애(충실성),
    추진성과에 의사결정 개선(효과성)과 이식성(확산가능성)을 같이 넣는다.

산출물
  ① SGIS활용사례_산불발화예측우선대응_하수범.docx        본문 5쪽
  ② 응모신청서_하수범.docx                                붙임1 신청서
  ③ 근거자료_하수범.docx                                  별도 첨부
"""

import os
import glob

import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

ROOT = os.path.join(r'C:', os.sep, 'for_sgis')
DERIVED = os.path.join(ROOT, 'data', 'grid_data', 'derived')
FIG = os.path.join(ROOT, 'outputs', 'figures')
DEST = os.path.join(r'C:', os.sep, 'Users', 'user', 'Desktop', '하수범_공모전', 'SGIS')
os.makedirs(DEST, exist_ok=True)

FONT = '맑은 고딕'
ACCENT = RGBColor(0x1F, 0x4E, 0x79)
TITLE = '산불 발화예측·우선대응 통합지도 — SGIS 통계지리 기반 500m 격자 의사결정 시스템'


def facts():
    f = {}
    r = pd.read_csv(os.path.join(DERIVED, 'ignition_ranks.csv'), encoding='utf-8-sig')
    b = r.groupby('fire_id')['haz_top_pct'].min().dropna()
    f['n_eval'] = len(b)
    for t in (1, 5, 10, 20):
        f[f'top{t}'] = (b <= t).mean() * 100
    da = r.groupby('fire_id')['damagearea'].first()
    d = pd.DataFrame({'best': b, 'ha': da}).dropna()
    f['med_big'] = d[d.ha >= 100].best.median()
    f['med_small'] = d[d.ha < 0.5].best.median()
    e = pd.read_parquet(os.path.join(DERIVED, 'mask_exposure_500m.parquet'))
    f['pop'] = e.pop_total.sum()
    v = pd.read_parquet(os.path.join(DERIVED, 'sgis_dong_vulnerability.parquet'))
    f['n_dong'] = len(v)
    ds = pd.read_csv(os.path.join(DERIVED, 'daily_scan_all.csv'), encoding='utf-8-sig')
    h = ds[ds.hour == 10]
    f['n_days'] = h.date.nunique()
    f['top1_day'] = h.top1_pop_day.mean()
    f['top1_res'] = h.top1_pop.mean()
    f['top5_old'] = h.top5_pop_old.mean()
    tops = [pd.read_csv(x, encoding='utf-8-sig')
            for x in sorted(glob.glob(os.path.join(DERIVED, 'replay_*_top.csv')))]
    a = pd.concat(tops, ignore_index=True)
    f['n_top'] = len(a)
    f['ratio_day'] = a.pop_day.sum() / a.pop_total.sum()
    f['n_case'] = len(tops)
    return f


F = facts()


# ── 문서 뼈대 ────────────────────────────────────────────────────────
def new_doc(body_pt=12.0, spacing=1.35):
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = FONT
    st.font.size = Pt(body_pt)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    st.paragraph_format.line_spacing = spacing
    st.paragraph_format.space_after = Pt(3)
    for s in doc.sections:
        # 안내문 지정 — 위·아래 10mm, 좌·우 20mm, 머리·꼬리말 10mm
        s.top_margin = s.bottom_margin = Cm(1.0)
        s.left_margin = s.right_margin = Cm(2.0)
        s.header_distance = s.footer_distance = Cm(1.0)
    return doc


def _run(p, text, size, bold=False, italic=False, color=None):
    r = p.add_run(text)
    r.font.name = FONT
    r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return r


def P(doc, text, size=12.0, bold=False, italic=False, align=None,
      before=0, after=3, indent=0.0, color=None):
    p = doc.add_paragraph()
    _run(p, text, size, bold, italic, color)
    if align:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    return p


def SEC(doc, text):
    """□ 로 시작하는 대절 — 양식이 지정한 4개 항목."""
    return P(doc, '□ ' + text, size=14.5, bold=True, before=11, after=5, color=ACCENT)


def ITEM(doc, text, size=12.0):
    return P(doc, '◦ ' + text, size=size, indent=0.3, after=2)


def SUB(doc, text, size=11.5):
    return P(doc, '- ' + text, size=size, indent=0.8, after=2)


def TABLE(doc, headers, rows, widths, size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        _run(c.paragraphs[0], h, size, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ''
            _run(cells[i].paragraphs[0], str(v), size)
    for row in t.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Cm(w)
        for c in row.cells:
            for p in c.paragraphs:
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.line_spacing = 1.0
    return t


def FIG_ROW(doc, names, caption, width):
    """그림을 한 줄에 나란히 — 5쪽 제약에서 세로 공간이 가장 비싸다."""
    t = doc.add_table(rows=1, cols=len(names))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, n in enumerate(names):
        p = t.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        f = os.path.join(FIG, f'{n}.png')
        if os.path.exists(f):
            p.add_run().add_picture(f, width=Cm(width))
    c = P(doc, caption, size=8.5, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    return c


# ══════════════════════════════════════════════════════════════════════
# ① 본문 — 5페이지 이내
# ══════════════════════════════════════════════════════════════════════
doc = new_doc()
P(doc, TITLE, size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=2, color=ACCENT)
P(doc, '“어디에 불이 날까”가 아니라 “어디를 먼저 지킬 것인가”에 답한다',
  size=11.5, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)

# ── 추진배경 및 필요성 ──
SEC(doc, '추진배경 및 필요성')
ITEM(doc, '산불 대응의 병목은 “불이 날까”가 아니라 “한정된 헬기와 진화대를 어디에 먼저 보낼 것인가”다.')
SUB(doc, '기존 산불위험지수는 시·군 단위 색상이라 그것만으로는 출동 순서를 정할 수 없다. '
          '같은 “위험” 등급이어도 사람이 사는 마을과 무인 산지는 지켜야 할 이유가 전혀 다르다.')
ITEM(doc, '2025년 경북 산불처럼 대형화가 반복되면서, 진화자원 배분의 우선순위가 곧 피해 규모를 가른다.')
SUB(doc, '발화 자체를 막을 수 없다면 최소한 “어디에 먼저 가 있을 것인가”는 미리 정할 수 있어야 한다.')
ITEM(doc, '위험도만으로는 부족하고 “그 격자에 무엇이 있는가”를 알아야 한다. 그 자리에 SGIS가 들어간다.')
SUB(doc, f'전국 500m 격자 403,385개의 발화 위험(t+1·2·3시간)에 SGIS 인구·가구·주택 격자통계와 '
          f'행정동 {F["n_dong"]:,}개의 인구특성을 결합해 대응 우선지역을 자동 산출한다.')

# ── 추진과정 ──
SEC(doc, '추진과정')
ITEM(doc, '활용분야 및 주요내용', size=11)
SUB(doc, '자료제공(500m 격자통계 인구·가구·주택), 개발지원센터 OpenAPI(지오코딩·행정동 경계·'
          '인구특성·통계주제도), 통계주제도(노후주택·보건시설 접근성)')
SUB(doc, f'2021~2025년 산불시즌 {F["n_days"]}일 전 기간을 매일 산출하고, 실제 발화 '
          f'{F["n_eval"]:,}건 전수로 검증했다.')
TABLE(doc, ['단계', '내용', 'SGIS의 역할'], [
    ['① 라벨 정비', '산불 이력을 격자·시간으로 정규화', '지오코딩 API로 좌표 결측 417건 복구'],
    ['② 격자 정합', '분석 격자와 SGIS 격자통계를 면적가중 결합', '격자경계 API로 원점 정렬 실측 검증'],
    ['③ 모델 학습', 'LightGBM 공간취약도 → GRU 시계열 2단', '(모델 단계는 SGIS 미사용)'],
    ['④ 노출 결합', '위험도 × 노출로 우선지역 산출', '주간인구 보정, 인구특성 결합'],
    ['⑤ 서비스화', '웹 지도·시간 슬라이더·분석 에이전트', '행정동 경계, 지역 공간질의'],
], widths=[2.1, 6.3, 8.6])

# ── 추진내용 ──
SEC(doc, '추진내용 — SGIS를 어디에 썼는가')
P(doc, '기준은 하나로 잡았다. SGIS를 빼면 결과가 달라지는가? 다섯 군데에서 달라진다.',
  size=11, bold=True, after=4)

ITEM(doc, '① 500m 격자 정합 — 면적가중 배분')
SUB(doc, f'분석 격자 한 셀이 SGIS 셀 4개에 걸치며 어긋난 양이 전 격자에서 같아 가중치가 상수다. '
          f'면적가중으로 배분해 총인구 {F["pop"]:,.0f}명(원본 대비 99.47%)을 실었다. '
          f'격자경계 API로 강릉·문경·강남의 꼭짓점을 실측해 500 배수 정렬을 확인했다.')

ITEM(doc, '② 지오코딩 복구 — 누락된 산불의 23%')
SUB(doc, '좌표 결측 417건을 SGIS 지오코딩 API 3단계 조회(리→읍면동→시군구)로 복구했다. '
          '여기에 2022년 울진 산불(16,302ha)을 비롯한 최대 규모 사건이 들어 있었고, '
          '복구 전 0.69였던 2022년 fold AUROC가 0.79로 회복됐다.')

ITEM(doc, '③ 노출의 시간 불일치 교정 — 핵심 기여')
SUB(doc, '산불은 오후에 몰려 스캔 시각을 11·14시로 잡았는데, 노출은 인구주택총조사 상주인구 '
          '곧 “밤에 어디서 자는가”로 재고 있었다. 낮에 예측하고 밤 인구로 피해를 셈한 셈이다.')
SUB(doc, 'SGIS에 주간인구·통근 통계가 없어 종사자 통계로 근사했다. '
          '주간지수 = (종사자 + 65세 이상 + 15세 미만) ÷ 상주인구.')
SUB(doc, f'그 결과 사례일 {F["n_case"]}일의 우선지역 Top10 중 56%가 교체됐고'
          f'({F["n_top"]:,}건 중 1,829건), 선정 격자의 주간인구는 상주인구의 {F["ratio_day"]:.2f}배가 됐다. '
          f'낮에 비는 곳 대신 낮에 차는 곳을 고른다.')
P(doc, '※ 한계 — 농작업은 사업체 등록에 잡히지 않는다. 논밭두렁 소각은 국내 산불 원인 1위인데 '
       '이 지표가 바로 그 활동을 세지 못해 농촌을 과소추정한다. 이 문구는 서비스 화면과 '
       '분석 에이전트에도 동일하게 표시한다.',
  size=9.5, indent=0.8, after=4, color=RGBColor(0x8B, 0x40, 0x00))

ITEM(doc, '④ 인구 특성 — 순위가 아니라 “무엇을 잃는가”')
SUB(doc, f'행정동 {F["n_dong"]:,}개의 연령대별 인구와 통계주제도 33종 중 30년 이상 노후주택·'
          f'보건시설 1개당 노인인구를 확보했다. 다만 우선순위 산식에는 넣지 않았다. '
          f'거주 격자끼리 비교하면 발화지의 고령비율은 0.91배, 노후주택비율은 0.89배로 오히려 낮았다.')
SUB(doc, '대신 “같은 위험도일 때 무엇을 잃는가”를 보여주는 표시용으로 쓴다. 화면은 '
          '“상위 5% 주간 노출인구” 아래에 65세 이상과 30년 이상 노후주택을 함께 세운다.')

ITEM(doc, '⑤ 행정경계와 공간질의')
SUB(doc, f'SGIS 행정동 경계 {F["n_dong"]:,}개를 지도에 얹어 지역을 검색·강조하고, 같은 코드 체계로 '
          f'분석 에이전트가 “의성군 상황은?”에 직접 답한다. 시도 코드가 SGIS 자체 체계라 '
          f'표준코드로는 이을 수 없어, 동 이름 집합을 상위 계층과 맞춰 252/252 완전일치로 복원했다.')

FIG_ROW(doc, ['03_daytime', '06_vulnerability'],
        '[그림 1] 주간인구 보정(좌) — 격자 대부분이 낮에 비워진다  |  '
        '[그림 2] 고령·노후주택은 발화와 무관해 산식에서 제외(우)', width=8.2)

# ── 추진성과 ──
SEC(doc, '추진성과')
ITEM(doc, f'실제 발화 {F["n_eval"]:,}건 전수 평가')
TABLE(doc, ['구간', '포착률', '무작위 대비', '비고'], [
    ['전국 상위 1%', f'{F["top1"]:.1f}%', f'{F["top1"]:.1f}배', '전국 4,034셀만 보고도'],
    ['전국 상위 5%', f'{F["top5"]:.1f}%', f'{F["top5"] / 5:.1f}배', ''],
    ['전국 상위 10%', f'{F["top10"]:.1f}%', f'{F["top10"] / 10:.1f}배', ''],
    ['100ha 이상 대형산불', f'중앙값 {F["med_big"]:.1f}%', '—',
     f'0.5ha 미만은 {F["med_small"]:.1f}% — 큰 불일수록 잘 잡는다'],
], widths=[3.6, 3.0, 2.4, 8.0])

ITEM(doc, '의사결정이 어떻게 달라지는가 (효과성)')
TABLE(doc, ['', '기존 방식', '본 시스템'], [
    ['공간 해상도', '시·군 단위 색상', '500m 격자 403,385개'],
    ['시간 해상도', '일 단위', 't+1h · t+2h · t+3h'],
    ['노출 기준', '없음 또는 상주인구(야간)', 'SGIS 주간 보정 인구'],
    ['산출물', '위험 등급', '대응 우선지역 순위 + 노출 규모 + 근거'],
], widths=[2.8, 5.6, 8.6])
SUB(doc, f'{F["n_days"]}일 평균으로 상위 1% 격자의 노출은 상주 {F["top1_res"]:,.0f}명에서 '
          f'주간 {F["top1_day"]:,.0f}명으로 바뀐다. 같은 위험지도라도 지켜야 할 대상이 달라진다. '
          f'그중 65세 이상은 평균 {F["top5_old"]:,.0f}명(상위 5% 기준)으로 함께 제시된다.')

ITEM(doc, '검증 태도 — 지표가 좋아진 모델을 기각했다 (창의성)')
SUB(doc, '신규 방법론(Stage2 CNN)으로 이관했다가 되돌렸다. 검증셋 AUROC는 0.837→0.884로 올랐는데 '
          '실제 발화 상위 1% 포착률은 5.3%→2.9%로 떨어졌기 때문이다(1,160건 전수, 5개 연도 전부 악화). '
          'AUROC는 40만 격자 순위 전체의 평균적 분리도이고 화면이 쓰는 “상위 1%”는 꼬리 4,034셀만 본다. '
          '서로 다른 것을 잰다.')

ITEM(doc, '“오늘은 위험한 날인가” — 시간축 등급')
SUB(doc, f'공간 백분위만 보면 조용한 날에도 상위 1%는 늘 나온다. {F["n_days"]}일 분포에서 오늘의 위치를 '
          f'따로 재는 지표를 두었다. 발화건수와의 스피어만 상관은 0.769이고, ‘매우 높음’ 등급인 날 중 '
          f'발화 0건인 날은 하나도 없었다.')

ITEM(doc, '왜 이 격자인가 — 설명 가능성')
SUB(doc, 'SHAP 대신 occlusion을 썼다. 입력이 12시간 시계열이라 “최근 몇 시간 중 어느 시점이 '
          '결정적이었나”가 곧 진화 지휘에 쓰이는 답이기 때문이다. 분석 결과 12시간을 넣지만 직전 '
          '1시간이 지배하고, 정적 입력에서는 Stage1 공간 취약도가 압도적이었다. 서비스에서는 '
          '우선지역 항목의 “왜?”를 누르면 이 기여도가 그대로 펼쳐진다.')

ITEM(doc, '확산 가능성')
SUB(doc, '격자 정합·노출 계산 모듈은 재난 종류와 무관하다. 위험도 레이어만 교체하면 호우·폭염·'
          '산사태에 그대로 쓸 수 있다. 주간인구 보정은 SGIS 종사자 통계만 있으면 어디서든 재현되며, '
          '낮에 일어나는 모든 재난에 공통으로 필요하다.')
SUB(doc, '“위험도 × 노출”을 두 백분위의 평균으로 정의한 방식은 단위가 다른 지표를 결합하는 '
          '일반적인 틀이라 다른 분야의 SGIS 활용에도 적용된다. 전 과정이 공개 저장소에 있어 '
          '지자체가 자기 관할로 좁혀 재현할 수 있다.')

FIG_ROW(doc, ['01_ignition_ranks', '04_ablation'],
        '[그림 3] 실제 발화 전수 평가(좌)  |  [그림 4] 검증셋 지표와 운영지표가 반대 방향(우)',
        width=8.2)
FIG_ROW(doc, ['02_time_axis', '05_occlusion'],
        '[그림 5] 시간축 위험등급별 실제 발화(좌)  |  [그림 6] occlusion 기여도 — 왜 이 격자인가(우)',
        width=8.2)

P(doc, 'SGIS를 빼면 이 시스템은 “위험한 곳”까지만 말하고 멈춘다. '
       '“먼저 지켜야 할 곳”은 SGIS가 있어야 나온다.',
  size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=6, color=ACCENT)

f1 = os.path.join(DEST, 'SGIS활용사례_산불발화예측우선대응_하수범.docx')
doc.save(f1)

# ══════════════════════════════════════════════════════════════════════
# ② 응모 신청서 (붙임1)
# ══════════════════════════════════════════════════════════════════════
d2 = new_doc(body_pt=10.5, spacing=1.2)
P(d2, 'SGIS 활용 우수사례 공모전 응모 신청서', size=16, bold=True,
  align=WD_ALIGN_PARAGRAPH.CENTER, after=10, color=ACCENT)

P(d2, '■ 인적 사항', size=12, bold=True, before=4, after=3)
TABLE(d2, ['항목', '내용'], [
    ['성명(팀명)', '하수범'],
    ['대표자 생년월일', '(작성 필요)'],
    ['대표 전화번호', '(작성 필요)'],
    ['대표 전자우편', 'chanvab1@gmail.com'],
    ['소속', '(작성 필요)'],
    ['공모전 유입 경로', '(작성 필요)'],
], widths=[4.5, 12.2], size=10)

P(d2, '■ 참가자 명단', size=12, bold=True, before=8, after=3)
TABLE(d2, ['성명', '생년월일', '연락처', '주소', '소속'], [
    ['하수범', '(작성 필요)', '(작성 필요)', '(작성 필요)', '(작성 필요)'],
], widths=[2.6, 3.0, 3.4, 4.6, 3.1], size=10)

P(d2, '■ SGIS 활용 분야', size=12, bold=True, before=8, after=3)
P(d2, '■ 자료제공        ■ 개발지원센터(OpenAPI)        ■ 통계주제도        '
      '■ 자연재해 통계지도', size=10.5, indent=0.3, after=2)
P(d2, '□ SGIS 전체  □ 대화형 통계지도  □ 생활권역통계지도  □ My 통계로  □ 일자리맵  '
      '□ 정책통계지도  □ 살고싶은 우리동네  □ 업종통계지도  □ 지역현안 소통지도  '
      '□ 기업생태 분석지도  □ 도시화 분석지도  □ 행정통계 시각화지도  □ 총조사 시각화지도  '
      '□ 월간통계  □ 인구피라미드  □ 고령화 현황보기  □ 성씨분포  □ 지방의 변화보기  '
      '□ SGIS 에듀  □ 기타', size=9, indent=0.3, after=6,
  color=RGBColor(0x80, 0x80, 0x80))

P(d2, '■ 작품 설명', size=12, bold=True, before=8, after=3)
TABLE(d2, ['구분', '내용'], [['제목', TITLE]], widths=[2.6, 14.1], size=10)
P(d2, '◦ (추진배경)', size=10.5, bold=True, indent=0.3, before=6, after=2)
P(d2, '- 산불 대응의 병목은 “불이 날까”가 아니라 “한정된 진화자원을 어디에 먼저 보낼 것인가”다. '
      '시·군 단위 위험등급만으로는 출동 순서를 정할 수 없어, 격자 단위 위험도에 '
      '“그 자리에 무엇이 있는가”를 결합할 필요가 있었다.', size=10, indent=0.8, after=4)
P(d2, '◦ (추진과정)', size=10.5, bold=True, indent=0.3, after=2)
P(d2, f'- SGIS 지오코딩 API로 산불 이력의 좌표 결측 417건을 복구하고, 격자경계 API로 정렬을 실측 검증한 뒤 '
      f'500m 격자통계를 면적가중 배분했다. 2단 모델(LightGBM+GRU)로 {F["n_days"]}일 전 기간을 매일 산출하고, '
      f'SGIS 인구특성으로 노출을 계산해 우선지역을 뽑았다.', size=10, indent=0.8, after=4)
P(d2, '◦ (주요내용)', size=10.5, bold=True, indent=0.3, after=2)
P(d2, f'- 전국 500m 격자 403,385개의 t+1~3시간 발화 위험을 산출하고, SGIS 인구·가구·주택 격자통계와 '
      f'행정동 {F["n_dong"]:,}개 인구특성을 결합해 대응 우선지역을 자동 선정한다. '
      f'특히 산불이 오후에 집중되는데 노출은 야간 상주인구로 재던 불일치를 SGIS 종사자 통계로 교정했다.',
  size=10, indent=0.8, after=4)
P(d2, '◦ (추진성과)', size=10.5, bold=True, indent=0.3, after=2)
P(d2, f'- 실제 발화 {F["n_eval"]:,}건 전수 평가에서 전국 상위 1%만 보고도 {F["top1"]:.1f}%를 포착했다'
      f'(무작위 대비 {F["top1"]:.1f}배). 주간인구 보정으로 우선지역 Top10의 56%가 교체됐고, '
      f'검증셋 지표가 좋아진 신규 모델을 운영지표로 재검증해 기각하는 등 검증 절차를 문서화했다. '
      f'전체 코드와 서비스를 공개했다.', size=10, indent=0.8, after=4)

f2 = os.path.join(DEST, '응모신청서_하수범.docx')
d2.save(f2)

# ══════════════════════════════════════════════════════════════════════
# ③ 근거자료 (별도 첨부 — 5쪽에 포함되지 않는다)
# ══════════════════════════════════════════════════════════════════════
d3 = new_doc(body_pt=10.5, spacing=1.25)
P(d3, '추진성과 근거자료', size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
  after=3, color=ACCENT)
P(d3, TITLE, size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)

P(d3, '1. 서비스 및 코드 공개', size=12.5, bold=True, before=6, after=3, color=ACCENT)
TABLE(d3, ['구분', '위치'], [
    ['서비스(데모)', 'https://wildfire-predict-framework.vercel.app'],
    ['전체 코드·데이터 처리', 'https://github.com/tradeprogram/wildfire_predict_framework'],
    ['방법론 상세', 'docs/METHODOLOGY.md · docs/ARCHITECTURE.md'],
    ['모델 검증 기록', 'docs/STAGE2_ABLATION.md'],
    ['그림 생성 스크립트', 'scripts/67_report_figures.py — 모든 수치를 산출물에서 직접 읽음'],
], widths=[4.2, 12.5], size=10)

P(d3, '2. 성능 근거', size=12.5, bold=True, before=10, after=3, color=ACCENT)
for n, cap in [
    ('01_ignition_ranks', f'[근거 1] 실제 발화 {F["n_eval"]:,}건 전수 평가 — 누적 포착률과 피해규모별 순위'),
    ('07_folds', '[근거 2] 연도별 leave-one-year-out 5-fold 성능'),
    ('02_time_axis', '[근거 3] 시간축 위험등급별 실제 발화 실적 — 스피어만 상관 0.769'),
]:
    p = d3.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f = os.path.join(FIG, f'{n}.png')
    if os.path.exists(f):
        p.add_run().add_picture(f, width=Cm(15.5))
    P(d3, cap, size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)

P(d3, '3. SGIS 결합 근거', size=12.5, bold=True, before=8, after=3, color=ACCENT)
for n, cap in [
    ('03_daytime', '[근거 4] SGIS 종사자 통계 기반 주간인구 보정'),
    ('06_vulnerability', '[근거 5] 고령·노후주택이 발화와 무관함 — 산식 제외의 근거'),
]:
    p = d3.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f = os.path.join(FIG, f'{n}.png')
    if os.path.exists(f):
        p.add_run().add_picture(f, width=Cm(15.5))
    P(d3, cap, size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)

P(d3, '4. 모델 검증 근거', size=12.5, bold=True, before=8, after=3, color=ACCENT)
for n, cap in [
    ('04_ablation', '[근거 6] 검증셋 지표가 좋아진 모델을 운영지표로 재검증해 기각'),
    ('05_occlusion', '[근거 7] occlusion 기여도 — 우선지역이 왜 위험한지의 설명'),
]:
    p = d3.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f = os.path.join(FIG, f'{n}.png')
    if os.path.exists(f):
        p.add_run().add_picture(f, width=Cm(15.5))
    P(d3, cap, size=9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)

P(d3, '5. 한계 (본문 분량 제약으로 여기에 정리)', size=12.5, bold=True,
  before=8, after=3, color=ACCENT)
for t in [
    '주간인구는 추정치다. 농작업을 세지 못해 농촌을 과소추정한다.',
    '소형 화재 식별력이 낮다. 0.5ha 미만은 순위 중앙값 30.2%다.',
    '하루 4개 시각만 스캔하므로 심야·저녁 발화 일부가 평가에서 빠진다.',
    '재표본화 학습이라 출력을 발생확률로 쓸 수 없다. 별도 보정이 필요하다.',
    '입력 기상 래스터 보유 범위 때문에 2025년 6월까지만 산출된다.',
    '전 기간 모드는 격자 단위 값을 저장하지 않아 공간질의가 시군구까지만 답한다.',
]:
    SUB(d3, t, size=10)
P(d3, '이 한계들은 서비스 화면과 분석 에이전트에도 같은 문구로 표시된다. '
      '보고서와 서비스가 다른 말을 하지 않게 하는 것이 이 프로젝트의 원칙이다.',
  size=10, italic=True, indent=0.3, before=4)

f3 = os.path.join(DEST, '근거자료_하수범.docx')
d3.save(f3)

for f in (f1, f2, f3):
    print(f'저장 {f}  ({os.path.getsize(f) / 1e3:.0f} KB)')
print(f'\n본문 문단 {len(doc.paragraphs)} · 표 {len(doc.tables)} · 그림 4장')
print(f'평가 사건 {F["n_eval"]:,}건 · 상위1% {F["top1"]:.1f}% · 전 기간 {F["n_days"]}일')
